import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional

import fitz  # PyMuPDF
import pdfplumber
import docx
import markdown

logger = logging.getLogger("rag_system.document_loader")

class DocumentLoader:
    """
    Service responsible for loading text from various document formats (PDF, DOCX, MD, TXT).
    Cleans text, preserves headings, splits by structural pages if applicable, and extracts metadata.
    """

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Removes repeated whitespaces and returns cleaned text, while preserving line endings
        associated with headings or paragraphs.
        """
        if not text:
            return ""
        # Replace multiple spaces/tabs with single space, keeping newlines
        cleaned = re.sub(r"[ \t]+", " ", text)
        # Collapse multiple newlines into a maximum of two newlines
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        return cleaned.strip()

    def load_pdf(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Loads a PDF file page by page using pdfplumber, falling back to PyMuPDF (fitz) if needed.
        """
        pages_content = []
        filename = file_path.name
        
        logger.info(f"Loading PDF file: {filename}")
        
        # Try pdfplumber first
        try:
            with pdfplumber.open(file_path) as pdf:
                if len(pdf.pages) == 0:
                    raise ValueError("PDF file is empty")
                
                for idx, page in enumerate(pdf.pages):
                    page_num = idx + 1
                    text = page.extract_text()
                    if text:
                        cleaned = self.clean_text(text)
                        if cleaned:
                            pages_content.append({
                                "text": cleaned,
                                "metadata": {
                                    "filename": filename,
                                    "page": page_num,
                                    "format": "pdf"
                                }
                            })
        except Exception as e:
            logger.warning(f"pdfplumber failed or encountered error reading {filename}: {str(e)}. Falling back to PyMuPDF.")
            pages_content = []
            
            # Fallback to PyMuPDF
            try:
                doc = fitz.open(file_path)
                if len(doc) == 0:
                    raise ValueError("PDF file is empty (PyMuPDF)")
                for idx in range(len(doc)):
                    page_num = idx + 1
                    page = doc.load_page(idx)
                    text = page.get_text()
                    if text:
                        cleaned = self.clean_text(text)
                        if cleaned:
                            pages_content.append({
                                "text": cleaned,
                                "metadata": {
                                    "filename": filename,
                                    "page": page_num,
                                    "format": "pdf"
                                }
                            })
            except Exception as fe:
                logger.error(f"PyMuPDF fallback failed for {filename}: {str(fe)}")
                raise IOError(f"Could not read PDF file {filename}: {str(fe)}") from fe
                
        if not pages_content:
            raise ValueError(f"No readable text could be extracted from PDF: {filename}")
            
        return pages_content

    def load_docx(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Loads a DOCX file using python-docx. Extracts paragraphs and tables,
        preserving headings when parsing. Since DOCX files don't have native fixed
        page concepts, we group content under logical headings or paginate by length.
        """
        filename = file_path.name
        logger.info(f"Loading DOCX file: {filename}")
        
        try:
            doc = docx.Document(file_path)
            paragraphs = doc.paragraphs
            
            if not paragraphs and not doc.tables:
                raise ValueError("DOCX file is empty")
                
            text_blocks = []
            current_section = "Introduction"
            
            # Simple heuristic: scan paragraphs, check styles
            for p in paragraphs:
                p_text = p.text.strip()
                if not p_text:
                    continue
                
                # Check for heading styles
                if p.style.name.startswith("Heading"):
                    current_section = p_text
                    # Keep the heading in the text flow, but marked
                    text_blocks.append(f"\n# {p_text}\n")
                else:
                    text_blocks.append(p_text)
                    
            # Extract tables if present and append to text blocks
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        table_text.append(" | ".join(row_data))
                if table_text:
                    text_blocks.append("\n" + "\n".join(table_text) + "\n")

            full_text = "\n".join(text_blocks)
            cleaned = self.clean_text(full_text)
            
            if not cleaned:
                raise ValueError("No text extracted from DOCX file")
                
            # Treat DOCX as 1 single logical page or split by ~1500 words per page to simulate pagination
            words = cleaned.split()
            words_per_page = 1000
            pages_content = []
            
            for i in range(0, len(words), words_per_page):
                page_num = (i // words_per_page) + 1
                page_words = words[i:i + words_per_page]
                page_text = " ".join(page_words)
                pages_content.append({
                    "text": page_text,
                    "metadata": {
                        "filename": filename,
                        "page": page_num,
                        "format": "docx"
                    }
                })
                
            return pages_content
        except Exception as e:
            logger.error(f"Error loading DOCX file {filename}: {str(e)}")
            raise IOError(f"Failed to read DOCX file {filename}: {str(e)}") from e

    def load_txt(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Loads a standard TXT file. Splits into pages based on standard word counts
        (1000 words per page).
        """
        filename = file_path.name
        logger.info(f"Loading TXT file: {filename}")
        
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
                
            cleaned = self.clean_text(content)
            if not cleaned:
                raise ValueError("TXT file is empty")
                
            words = cleaned.split()
            words_per_page = 1000
            pages_content = []
            
            for i in range(0, len(words), words_per_page):
                page_num = (i // words_per_page) + 1
                page_words = words[i:i + words_per_page]
                page_text = " ".join(page_words)
                pages_content.append({
                    "text": page_text,
                    "metadata": {
                        "filename": filename,
                        "page": page_num,
                        "format": "txt"
                    }
                })
            return pages_content
        except Exception as e:
            logger.error(f"Error loading TXT file {filename}: {str(e)}")
            raise IOError(f"Failed to read TXT file {filename}: {str(e)}") from e

    def load_markdown(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Loads a Markdown file. Similar to TXT, but retains markdown heading formatting.
        """
        filename = file_path.name
        logger.info(f"Loading Markdown file: {filename}")
        
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
                
            cleaned = self.clean_text(content)
            if not cleaned:
                raise ValueError("Markdown file is empty")
                
            # Check if there are headings to determine structure
            # For simplicity of pagination, we split by ~1000 words
            words = cleaned.split()
            words_per_page = 1000
            pages_content = []
            
            for i in range(0, len(words), words_per_page):
                page_num = (i // words_per_page) + 1
                page_words = words[i:i + words_per_page]
                page_text = " ".join(page_words)
                pages_content.append({
                    "text": page_text,
                    "metadata": {
                        "filename": filename,
                        "page": page_num,
                        "format": "md"
                    }
                })
            return pages_content
        except Exception as e:
            logger.error(f"Error loading Markdown file {filename}: {str(e)}")
            raise IOError(f"Failed to read Markdown file {filename}: {str(e)}") from e

    def load_document(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Dispatches loading based on file extension.
        """
        ext = file_path.suffix.lower()
        if ext == ".pdf":
            return self.load_pdf(file_path)
        elif ext == ".docx":
            return self.load_docx(file_path)
        elif ext == ".txt":
            return self.load_txt(file_path)
        elif ext in (".md", ".markdown"):
            return self.load_markdown(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
