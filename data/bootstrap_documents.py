import os
import sys
from pathlib import Path

# Add project root to path to ensure dependencies are importable if needed
sys.path.append(str(Path(__file__).resolve().parent.parent))

def create_python_doc(dest_dir: Path):
    file_path = dest_dir / "python_documentation.txt"
    content = """PYTHON DOCUMENTATION - ADVANCED FEATURES
======================================

Section 1: Asynchronous Programming
-----------------------------------
Asynchronous programming in Python is supported natively by the asyncio library. It allows writing concurrent code using the async and await syntax. Asyncio is used as a foundation for multiple Python asynchronous frameworks that provide high-performance network and web-servers, database connection libraries, and distributed task queues.
The event loop is the core of every asyncio application. Event loops run asynchronous tasks and callbacks, perform network I/O operations, and run sub-processes.
An async function is defined with "async def". Inside async functions, you can use "await" to pause execution until the awaited coroutine finishes, allowing other operations to run concurrently.

Section 2: List Comprehensions
------------------------------
List comprehensions provide a concise way to create lists. Common applications are to make new lists where each element is the result of some operations applied to each member of another sequence or iterable, or to create a subsequence of those elements that satisfy a certain condition.
Syntax: [expression for item in iterable if condition]
This syntax is more readable and often faster than standard for-loops because it is optimized in C under the hood.

Section 3: Generators and Iterators
-----------------------------------
A generator is a special type of iterator that yields values on demand using the 'yield' keyword instead of returning them all at once. This makes generators highly memory-efficient when dealing with large datasets or infinite streams.
"""
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"Created TXT: {file_path}")

def create_fastapi_guide(dest_dir: Path):
    file_path = dest_dir / "fastapi_guide.md"
    content = """# FastAPI Guide: Modern Web Applications

## Introduction to FastAPI
FastAPI is a modern, fast (high-performance), web framework for building APIs with Python 3.8+ based on standard Python type hints.
Key features include:
* **Fast**: Very high performance, on par with NodeJS and Go (thanks to Starlette and Uvicorn).
* **Fast to code**: Increase the speed to develop features by about 200% to 300%.
* **Fewer bugs**: Reduce about 40% of developer induced errors.
* **Intuitive**: Great editor support. Completion everywhere. Less time debugging.

## Path Parameters and Validation
FastAPI allows you to declare path parameters and validate them using standard Python type hints.
```python
from fastapi import FastAPI

app = FastAPI()

@app.get("/items/{item_id}")
def read_item(item_id: int):
    return {"item_id": item_id}
```
If a user requests `/items/foo`, FastAPI will automatically return a detailed 422 validation error stating that `item_id` must be an integer, rather than throwing a internal server crash.

## Dependency Injection
FastAPI has a very powerful and intuitive Dependency Injection system. It allows you to share database connections, enforce security policies, retrieve settings, and load components cleanly.
You declare dependencies using the `Depends` class.
"""
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"Created MD: {file_path}")

def create_ml_notes(dest_dir: Path):
    file_path = dest_dir / "machine_learning_notes.txt"
    content = """MACHINE LEARNING ESSENTIAL NOTES
================================

Section 1: Supervised Learning
------------------------------
Supervised learning is the machine learning task of learning a function that maps an input to an output based on example input-output pairs. It infers a function from labeled training data consisting of a set of training examples.
Common algorithms:
1. Linear Regression (for continuous values)
2. Logistic Regression (for binary classification)
3. Support Vector Machines (SVM)
4. Random Forests and Decision Trees

Section 2: Gradient Descent
---------------------------
Gradient descent is an optimization algorithm used to minimize some cost function by iteratively moving in the direction of steepest descent as defined by the negative of the gradient. In machine learning, we use gradient descent to update the parameters (weights and biases) of our model.
The size of steps taken to reach the minimum is determined by the learning rate. If the learning rate is too high, the algorithm may overshoot the minimum and diverge. If it is too low, the algorithm will take too long to converge.

Section 3: Bias-Variance Tradeoff
---------------------------------
- Bias: Error introduced by approximating a real-world problem (which may be extremely complex) by a much simpler model. High bias leads to underfitting.
- Variance: Error introduced by the model's sensitivity to small fluctuations in the training set. High variance leads to overfitting.
The tradeoff is the sweet spot where the combined error of bias and variance is minimized.
"""
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"Created TXT: {file_path}")

def create_hr_policy(dest_dir: Path):
    file_path = dest_dir / "company_hr_policy.docx"
    import docx
    
    doc = docx.Document()
    doc.add_heading("Company Human Resources Policy Handbook", 0)
    
    doc.add_heading("Section 1: Leave and Paid Time Off (PTO)", 1)
    doc.add_paragraph(
        "All full-time employees are entitled to 20 days of paid annual leave per calendar year. "
        "PTO requests must be submitted through the HR Portal at least two weeks in advance. "
        "Unused annual leave up to a maximum of 5 days can be carried forward to the next year. "
        "Maternity leave is provided up to 26 weeks of paid leave, and paternity leave is provided up to 4 weeks of paid leave."
    )
    
    doc.add_heading("Section 2: Remote Work Policy", 1)
    doc.add_paragraph(
        "Employees are eligible for a hybrid working model, permitting remote work for up to 3 days per week, "
        "subject to approval by the direct supervisor. Core collaboration hours are between 10:00 AM and 3:00 PM EST. "
        "The company provides a one-time home office equipment stipend of $500 to assist with setup expenses."
    )

    doc.add_heading("Section 3: Equipment Allowance and Tech Stipend", 1)
    doc.add_paragraph(
        "Full-time employees receive a company-issued laptop which remains the property of the company. "
        "A monthly technology stipend of $50 is added to payroll to offset home internet and phone bills."
    )
    
    doc.save(file_path)
    print(f"Created DOCX: {file_path}")

def create_cybersecurity_handbook(dest_dir: Path):
    file_path = dest_dir / "cybersecurity_handbook.pdf"
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    
    doc = SimpleDocTemplate(str(file_path), pagesize=letter)
    story = []
    
    styles = getSampleStyleSheet()
    
    # Custom heading styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Title'],
        fontName='Helvetica-Bold',
        fontSize=22,
        spaceAfter=15
    )
    h1_style = ParagraphStyle(
        'DocH1',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=15,
        spaceBefore=10,
        spaceAfter=10
    )
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        spaceAfter=8
    )

    story.append(Paragraph("Enterprise Cybersecurity Policy Handbook", title_style))
    story.append(Spacer(1, 15))
    
    story.append(Paragraph("Section 1: Password Complexity and Authentication", h1_style))
    story.append(Paragraph(
        "All corporate passwords must be at least 14 characters in length and include a mix of "
        "uppercase letters, lowercase letters, numbers, and special characters (e.g., @, #, $). "
        "Passwords must be updated every 90 days. Sharing passwords or credentials is strictly prohibited. "
        "Multi-Factor Authentication (MFA) must be enabled on all systems containing corporate data.",
        body_style
    ))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Section 2: Phishing Prevention and Email Security", h1_style))
    story.append(Paragraph(
        "Employees must exercise caution when opening email attachments or clicking links in messages "
        "from external senders. Always inspect the sender email domain for discrepancies. "
        "Report suspicious emails immediately using the 'Report Phishing' button in the mail client. "
        "The IT department conducts random simulated phishing tests throughout the year.",
        body_style
    ))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Section 3: Incident Response and Reporting Procedures", h1_style))
    story.append(Paragraph(
        "In the event of a suspected security breach, data leak, or lost corporate device, "
        "the employee must report the incident to the Security Operations Center (SOC) within 1 hour. "
        "Do not attempt to investigate or clean the device yourself. The SOC is reachable 24/7 "
        "at soc@company.com or via extension 5555.",
        body_style
    ))

    doc.build(story)
    print(f"Created PDF: {file_path}")

def main():
    dest_dir = Path(__file__).resolve().parent / "documents"
    dest_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"Bootstrapping mock documents into: {dest_dir}")
    create_python_doc(dest_dir)
    create_fastapi_guide(dest_dir)
    create_ml_notes(dest_dir)
    
    # Try importing docx and reportlab to generate binary files
    try:
        import docx
        create_hr_policy(dest_dir)
    except ImportError:
        print("python-docx is not installed. Skipping DOCX generation. Run 'pip install python-docx'")
        
    try:
        import reportlab
        create_cybersecurity_handbook(dest_dir)
    except ImportError:
        print("reportlab is not installed. Skipping PDF generation. Run 'pip install reportlab'")

    # Create a gitkeep file for vector_db directory to ensure folders exist
    db_dir = Path(__file__).resolve().parent / "vector_db"
    db_dir.mkdir(parents=True, exist_ok=True)
    with open(db_dir / ".gitkeep", "w") as f:
        f.write("")

if __name__ == "__main__":
    main()
